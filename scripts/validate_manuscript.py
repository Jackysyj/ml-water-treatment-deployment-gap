#!/usr/bin/env python3
"""
Manuscript validator for the npj Clean Water Perspective draft.

Single sources of truth:
  - data/literature/extracted/fulltext_extraction_clean.json
  - figures/manuscript_stats.json
  - figures/plant_panel_stats.json
  - figures/verification_summary.json
  - figures/fig1_data_summary.csv
  - figures/fig2_operational_illustration_summary.csv
"""

from __future__ import annotations

import csv
import json
import re
import string
import sys
from zipfile import ZipFile
from pathlib import Path

import pandas as pd
from docx import Document

from paths import ANALYSIS_CORPUS_JSON, FIGURES_DIR, DRAFTS_DIR, SUBMISSION_DIR


DRAFTS = DRAFTS_DIR
FIGURES = FIGURES_DIR
SUBMISSION = SUBMISSION_DIR
DATA_PATH = ANALYSIS_CORPUS_JSON

TITLE = "Operational Evidence Standards for Machine Learning in Wastewater Treatment"
MAX_TITLE_WORDS = 15
MAX_ABSTRACT_WORDS = 70
MAX_MAIN_WORDS = 3000
MAX_REFERENCES = 70
MAX_FIGURE_LEGEND_WORDS = 350
MIN_COVER_LETTER_WORDS = 560
MAX_COVER_LETTER_WORDS = 760
CITATION_BLUE = "0563C1"
CITATION_MARKER_RE = re.compile(r"\[(?:\d+(?:\s*-\s*\d+)?)(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*\]")
MAIN_TEXT_FILES = [
    "introduction.md",
    "sec1_landscape.md",
    "sec2_data_quality.md",
    "sec3_it_ot.md",
    "sec4_trust.md",
    "sec6_discussion.md",
]
BODY_FILES = ["abstract.md", *MAIN_TEXT_FILES, "figure_captions.md"]
ALL_PUBLIC_FILES = [
    *BODY_FILES,
    "data_availability.md",
    "code_availability.md",
    "acknowledgements.md",
    "author_contributions.md",
    "competing_interests.md",
    "cover_letter.md",
    "highlights.md",
    "SI_supporting_information.md",
]

GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
RESET = "\033[0m"


def load_json(path: Path):
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def load_text(name: str) -> str:
    return (DRAFTS / name).read_text(encoding="utf-8")


def strip_heading(text: str) -> str:
    return "\n".join(ln for ln in text.splitlines() if not ln.strip().startswith("#")).strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def passfail(condition: bool, scope: str, check: str, detail: str = ""):
    return ("PASS" if condition else "FAIL", scope, check, detail)


def warn(condition: bool, scope: str, check: str, detail: str = ""):
    return ("PASS" if condition else "WARN", scope, check, detail)


def load_stats():
    raw = load_json(DATA_PATH)
    df = pd.DataFrame(raw["results"])
    ms = load_json(FIGURES / "manuscript_stats.json")
    plant = load_json(FIGURES / "plant_panel_stats.json")
    ver = load_json(FIGURES / "verification_summary.json")
    return df, ms, plant, ver


def load_main_text() -> str:
    return "\n".join(load_text(name) for name in MAIN_TEXT_FILES)


def load_body_text() -> str:
    return "\n".join(load_text(name) for name in BODY_FILES)


def check_perspective_format():
    checks = []
    punctuation = set(string.punctuation) - {"-"}
    title_words = word_count(TITLE)
    checks.append(passfail(title_words <= MAX_TITLE_WORDS, "format", "title <=15 words", str(title_words)))
    checks.append(passfail(not any(ch in punctuation for ch in TITLE), "format", "title has no punctuation", TITLE))

    abstract_words = word_count(strip_heading(load_text("abstract.md")))
    checks.append(passfail(abstract_words <= MAX_ABSTRACT_WORDS, "format", "abstract <=70 words", str(abstract_words)))

    main_words = word_count(load_main_text())
    checks.append(passfail(main_words <= MAX_MAIN_WORDS, "format", "main text <=3000 words", str(main_words)))

    captions = re.split(r"\n## Figure \d+\n", load_text("figure_captions.md"))[1:]
    checks.append(passfail(len(captions) == 3, "format", "three main figure legends", str(len(captions))))
    for i, caption in enumerate(captions, start=1):
        checks.append(passfail(word_count(caption) <= MAX_FIGURE_LEGEND_WORDS, "format", f"Fig. {i} legend <=350 words", str(word_count(caption))))

    n_refs = len(re.findall(r"^\[\d+\]", load_text("ref.md"), flags=re.M))
    checks.append(passfail(n_refs <= MAX_REFERENCES, "format", "references <=70", str(n_refs)))
    return checks


def check_truth_numbers(df: pd.DataFrame, ms: dict, plant: dict, ver: dict):
    checks = []
    checks.append(passfail(len(df) == 423, "truth", "analysis corpus n=423", str(len(df))))
    checks.append(passfail(ms.get("total_papers") == 423, "truth", "manuscript_stats total", str(ms.get("total_papers"))))
    expected = {
        "r2_n": 202,
        "r2_median": 0.97,
        "deployed_n": 12,
        "deployed_pct": 2.8,
        "realtime_pct": 5.2,
        "uq_pct": 8.5,
        "code_pct": 7.3,
        "data_avail_pct": 13.2,
    }
    for key, expected_value in expected.items():
        checks.append(passfail(ms.get(key) == expected_value, "truth", key, str(ms.get(key))))

    dep = ver["deployment_verification"]
    val = ver["validation_verification"]
    checks.append(passfail(dep["n_deployed_labels_in_corpus"] == 12, "truth", "12 deployment labels audited", str(dep)))
    checks.append(passfail(dep["n_confirmed_full_text"] == 12, "truth", "12 deployment labels confirmed", str(dep)))
    checks.append(passfail(val["n_rigorous_labels_in_corpus"] == 77, "truth", "77 future-facing validation labels audited", str(val)))

    total_records = sum(p["n_complete"] for p in plant["plants"])
    predictors = [p["n_predictors"] for p in plant["plants"]]
    checks.append(passfail(plant["n_plants"] == 6, "truth", "six plants", str(plant["n_plants"])))
    checks.append(passfail(total_records == 6006, "truth", "6,006 plant records", str(total_records)))
    checks.append(passfail(min(predictors) == 8 and max(predictors) == 13, "truth", "8 to 13 predictors", str(predictors)))
    scenarios = plant.get("scenario_summary", {})
    expected_scenarios = {
        "random_cv_influent_only": 0.063,
        "walk_forward_mean_baseline": -0.024,
        "walk_forward_persistence": -0.960,
        "walk_forward_influent_only": 0.029,
        "walk_forward_influent_plus_lagged_effluent": 0.031,
    }
    for key, expected_median in expected_scenarios.items():
        checks.append(passfail(key in scenarios, "truth", f"scenario present: {key}", ""))
        if key in scenarios:
            checks.append(passfail(round(float(scenarios[key]["median"]), 3) == expected_median, "truth", f"{key} median", str(scenarios[key])))
            checks.append(passfail(scenarios[key].get("n") == 29, "truth", f"{key} n=29", str(scenarios[key])))
    return checks


def check_text_sync():
    checks = []
    required = {
        "abstract.md": [r"423 full-text", r"2\.8%", r"5\.2%", r"18\.2%", r"8\.5%", r"six-plant"],
        "introduction.md": [r"median reported \*R²\* of 0\.97", r"423-study full-text evidence map", r"all 12 plant-deployment labels", r"29 plant-target tasks"],
        "sec1_landscape.md": [r"423-study full-text corpus", r"68\.6%", r"290 of 423", r"202 studies", r"median reported value is 0\.97"],
        "sec2_data_quality.md": [r"12 of 423", r"2\.8%", r"5\.2%", r"18\.2%", r"8\.5%", r"7\.3%", r"13\.2%", r"deployed events total five"],
        "sec3_it_ot.md": [r"6,006 timestamped records", r"8 to 13", r"0\.063", r"0\.029", r"-0\.024", r"-0\.960", r"0\.031"],
        "sec4_trust.md": [r"operation-aware evidence standard", r"47 of 423", r"11\.1%", r"17\.4%", r"shadow mode", r"advisory mode"],
        "data_availability.md": [r"423-study full-text analysis corpus", r"Raw plant records are not publicly available", r"derived outputs"],
        "code_availability.md": [r"analysis, validation and figure-generation scripts"],
    }
    for filename, patterns in required.items():
        text = load_text(filename)
        for pattern in patterns:
            checks.append(passfail(re.search(pattern, text) is not None, "text sync", f"{filename}: {pattern}", ""))
    table_links = {
        "sec2_data_quality.md": ["Supplementary Table S11", "Supplementary Table S12"],
        "sec3_it_ot.md": ["Supplementary Table S13a", "Supplementary Table S13b"],
    }
    for filename, links in table_links.items():
        text = load_text(filename)
        for link in links:
            checks.append(passfail(link in text, "text sync", f"{filename}: cites {link}", ""))
    return checks


def check_figure_outputs_and_links():
    checks = []
    expected = {
        1: "fig1_landscape.png",
        2: "fig2_operational_illustration.png",
        3: "fig3_evidence_framework.png",
    }
    for fig_num, filename in expected.items():
        checks.append(passfail((FIGURES / filename).exists(), "figures", f"Fig. {fig_num} file exists", filename))
    main = load_main_text()
    cited = set(int(m.group(1)) for m in re.finditer(r"\b(?:Fig\.|Figure)\s+([1-9])", main))
    checks.append(passfail(cited == {1, 2, 3}, "figures", "main text cites only Fig. 1-3", str(sorted(cited))))

    captions = load_text("figure_captions.md")
    checks.append(passfail(re.search(r"(?m)^\(a\) ", captions) is not None, "figures", "caption panel labels use (a)", ""))
    checks.append(passfail(re.search(r"(?m)^[a-z]\) ", captions) is None, "figures", "caption panel labels do not use a)", ""))
    checks.append(passfail("## Figure 4" not in captions and "## Figure 5" not in captions, "figures", "no Fig. 4 or Fig. 5 caption", ""))
    checks.append(passfail((FIGURES / "fig_s1_prisma.png").exists(), "figures", "SI Fig. S1 file exists", "fig_s1_prisma.png"))
    checks.append(passfail((FIGURES / "fig_s2_subfield_distribution.png").exists(), "figures", "SI Fig. S2 file exists", "fig_s2_subfield_distribution.png"))
    return checks


def check_guardrails():
    checks = []
    text = "\n".join(load_text(name) for name in ALL_PUBLIC_FILES if (DRAFTS / name).exists())
    forbidden = {
        "Nature Water": "old target journal in submission-facing text",
        "Research Article": "Article type in submission-facing text",
        "does not track": "underpowered association wording",
        "decoupled": "underpowered association wording",
        "none approaches literature median": "plant-vs-literature comparison",
        "literature median compared": "plant-vs-literature comparison",
        "against the literature median": "plant-vs-literature comparison",
        "unrecoverable": "over-strong sensing-boundary wording",
        "single plant": "old single-plant framing",
        "1,682": "single-plant sample headline",
        "readiness score": "composite score wording",
    }
    for token, desc in forbidden.items():
        checks.append(passfail(token not in text, "guardrails", desc, token))

    body = load_body_text()
    checks.append(passfail("—" not in body, "guardrails", "no em dash in manuscript body", ""))
    checks.append(passfail("[ref]" not in body.lower() and "[cite]" not in body.lower(), "guardrails", "no citation placeholders in body", ""))
    checks.append(passfail("[Author names]" not in text and "[Affiliations]" not in text and "[Month Year]" not in text, "guardrails", "no title-page placeholders", ""))
    checks.append(passfail("to be added before submission" not in text and "to be completed before submission" not in text, "guardrails", "no administrative placeholders", ""))
    checks.append(passfail("inflated" not in body.lower(), "guardrails", "no inflated-accuracy thesis", ""))
    checks.append(passfail("0.97" not in load_text("sec3_it_ot.md"), "guardrails", "6-plant section has no 0.97 comparison", ""))
    return checks


def expand_reference_group(group: str) -> list[int]:
    refs = []
    for part in group.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start, end = part.split("-", 1)
            if start.strip().isdigit() and end.strip().isdigit():
                refs.extend(range(int(start), int(end) + 1))
        elif part.isdigit():
            refs.append(int(part))
    return refs


def check_reference_integrity():
    checks = []
    main = load_main_text()
    ref_text = load_text("ref.md")
    seq = []
    groups = []
    for match in re.finditer(r"\[([0-9][0-9,\-\s]*)\]", main):
        refs = expand_reference_group(match.group(1))
        if refs:
            seq.extend(refs)
            groups.append(refs)
    used = sorted(set(seq))
    defined = sorted(int(n) for n in re.findall(r"^\[(\d+)\]", ref_text, flags=re.M))
    first_order = []
    seen = set()
    for ref in seq:
        if ref not in seen:
            first_order.append(ref)
            seen.add(ref)
    checks.append(passfail(defined == list(range(1, len(defined) + 1)), "references", "reference list numbered consecutively", str(defined)))
    checks.append(passfail(not (set(used) - set(defined)), "references", "all cited references are defined", str(sorted(set(used) - set(defined)))))
    checks.append(passfail(not (set(defined) - set(used)), "references", "all defined references are cited", str(sorted(set(defined) - set(used)))))
    checks.append(passfail(first_order == sorted(first_order), "references", "first citation order is monotonic", str(first_order)))
    checks.append(passfail(all(group == sorted(group) for group in groups), "references", "citation groups are ascending", str(groups)))
    return checks


def check_si_closure():
    checks = []
    main = load_main_text()
    si = load_text("SI_supporting_information.md")
    checks.append(passfail("Operational Evidence Standards for Machine Learning in Wastewater Treatment" in si, "SI", "SI title updated", ""))
    checks.append(passfail("main Figure 4" not in si and "main Figure 5" not in si, "SI", "no old main Fig. 4/5 references", ""))
    checks.append(passfail("All 500 papers in the final corpus" not in si, "SI", "no all-500 final-corpus wording", ""))
    for table in ["S11", "S12", "S13a", "S13b", "S13c", "S14"]:
        checks.append(passfail(f"Table {table}." in si, "SI", f"Table {table} exists", ""))
    checks.append(passfail("[See `figures/" not in si, "SI", "no bare SI figure placeholders in source", ""))
    return checks


def _docx_text(path: Path) -> str:
    if not path.exists():
        return ""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _docx_has_page_field(path: Path) -> bool:
    if not path.exists():
        return False
    with ZipFile(path) as zf:
        footer_names = [name for name in zf.namelist() if name.startswith("word/footer")]
        return any(" PAGE " in zf.read(name).decode("utf-8", errors="ignore") for name in footer_names)


def _docx_image_count(path: Path) -> int:
    if not path.exists():
        return -1
    return len(Document(path).inline_shapes)


def _run_rgb(run):
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    return str(color.rgb).upper()


def _docx_citation_colour_stats(path: Path) -> dict:
    stats = {
        "citation_runs_before_refs": 0,
        "blue_citation_runs_before_refs": 0,
        "nonblue_citation_runs_before_refs": [],
        "blue_noncitation_runs_before_refs": [],
        "blue_reference_number_runs_after_refs": [],
    }
    if not path.exists():
        return stats

    in_references = False
    doc = Document(path)
    for para in doc.paragraphs:
        if para.text.strip() == "References":
            in_references = True
            continue
        for run in para.runs:
            text = run.text
            if not text:
                continue
            rgb = _run_rgb(run)
            is_blue = rgb == CITATION_BLUE
            markers = CITATION_MARKER_RE.findall(text)
            is_marker_only = CITATION_MARKER_RE.fullmatch(text.strip()) is not None

            if not in_references:
                if markers:
                    if is_marker_only:
                        stats["citation_runs_before_refs"] += 1
                        if is_blue:
                            stats["blue_citation_runs_before_refs"] += 1
                        else:
                            stats["nonblue_citation_runs_before_refs"].append(text)
                    else:
                        stats["nonblue_citation_runs_before_refs"].append(text)
                elif is_blue and text.strip():
                    stats["blue_noncitation_runs_before_refs"].append(text)
            elif is_blue and re.match(r"^\[\d+\]", text.strip()):
                stats["blue_reference_number_runs_after_refs"].append(text)
    return stats


def check_submission_package():
    checks = []
    main_docx = SUBMISSION / "main_manuscript.docx"
    si_docx = SUBMISSION / "supporting_information.docx"
    cover_docx = SUBMISSION / "cover_letter.docx"
    highlights = SUBMISSION / "highlights.txt"

    for path in [main_docx, si_docx, cover_docx, highlights]:
        checks.append(passfail(path.exists(), "submission", f"{path.name} exists", str(path)))

    main_text = _docx_text(main_docx)
    si_text = _docx_text(si_docx)
    cover_text = _docx_text(cover_docx)

    checks.append(passfail(TITLE in main_text, "submission", "main DOCX title page has title", ""))
    checks.append(passfail("Siyuan Jiang, Ying Yang, Keqiang Zhang, Yaoru Mao, Xiuwen Cheng*" in main_text, "submission", "main DOCX title page has authors", ""))
    checks.append(passfail("Submitted to npj Clean Water, June 2026" in main_text, "submission", "main DOCX submitted line", ""))
    checks.append(passfail(_docx_has_page_field(main_docx), "submission", "main DOCX has page-number field", ""))
    checks.append(passfail(_docx_image_count(main_docx) == 3, "submission", "main DOCX has 3 embedded figures", str(_docx_image_count(main_docx))))
    checks.append(passfail("Figure and Table Legends" not in main_text, "submission", "main figures not collected under detached legend heading", ""))
    citation_stats = _docx_citation_colour_stats(main_docx)
    checks.append(passfail(citation_stats["citation_runs_before_refs"] > 0, "submission", "main DOCX has citation-marker runs", str(citation_stats["citation_runs_before_refs"])))
    checks.append(passfail(
        citation_stats["citation_runs_before_refs"] == citation_stats["blue_citation_runs_before_refs"],
        "submission",
        "main-text citation markers are hyperlink blue",
        str(citation_stats),
    ))
    checks.append(passfail(not citation_stats["blue_noncitation_runs_before_refs"], "submission", "non-citation main text is not blue", str(citation_stats["blue_noncitation_runs_before_refs"][:5])))
    checks.append(passfail(not citation_stats["blue_reference_number_runs_after_refs"], "submission", "reference-list numbers are not blue", str(citation_stats["blue_reference_number_runs_after_refs"][:5])))

    checks.append(passfail("Supporting Information" in si_text, "submission", "SI DOCX title page exists", ""))
    checks.append(passfail(TITLE in si_text, "submission", "SI DOCX repeats manuscript title", ""))
    checks.append(passfail("Contents" in si_text and "S1. Literature Search" in si_text, "submission", "SI DOCX Contents exists", ""))
    checks.append(passfail(_docx_has_page_field(si_docx), "submission", "SI DOCX has page-number field", ""))
    checks.append(passfail(_docx_image_count(si_docx) == 2, "submission", "SI DOCX has 2 embedded figures", str(_docx_image_count(si_docx))))
    checks.append(passfail("[See `figures/" not in si_text and "![Figure" not in si_text, "submission", "SI DOCX has no figure placeholders", ""))

    placeholder_re = re.compile(r"\[[A-Za-z][^\]]+\]|to be added before submission|to be completed before submission", re.I)
    checks.append(passfail(not placeholder_re.search(main_text), "submission", "main DOCX has no administrative placeholders", ""))
    checks.append(passfail(not placeholder_re.search(si_text), "submission", "SI DOCX has no administrative placeholders", ""))
    checks.append(passfail(not placeholder_re.search(cover_text), "submission", "cover letter has no administrative placeholders", ""))
    checks.append(passfail(TITLE in cover_text and "Perspective" in cover_text and "chengxw@lzu.edu.cn" in cover_text, "submission", "cover letter contains title, type and corresponding email", ""))
    cover_words = word_count(cover_text)
    checks.append(passfail(MIN_COVER_LETTER_WORDS <= cover_words <= MAX_COVER_LETTER_WORDS, "submission", "cover letter targets two-page length", str(cover_words)))

    if highlights.exists():
        lines = [line.strip() for line in highlights.read_text(encoding="utf-8").splitlines() if line.strip()]
        stripped = [re.sub(r"^[-*]\s*", "", line) for line in lines]
        checks.append(passfail(3 <= len(stripped) <= 5, "submission", "highlights has 3-5 items", str(len(stripped))))
        checks.append(passfail(all(len(line) <= 85 for line in stripped), "submission", "each highlight <=85 characters", str([len(line) for line in stripped])))
    return checks


def print_section(title: str, checks):
    print(f"\n{title}")
    passed = failed = warned = 0
    for status, scope, check, detail in checks:
        color = GREEN if status == "PASS" else RED if status == "FAIL" else YELLOW
        print(f"  {color}{status:<5}{RESET}  {scope:<12}  {check:<58}  {detail}")
        if status == "PASS":
            passed += 1
        elif status == "FAIL":
            failed += 1
        else:
            warned += 1
    return passed, failed, warned


def main():
    print("=" * 72)
    print("  npj Clean Water Perspective validator")
    print("=" * 72)
    df, ms, plant, ver = load_stats()
    sections = [
        ("1. Perspective format", check_perspective_format()),
        ("2. Truth numbers", check_truth_numbers(df, ms, plant, ver)),
        ("3. Text/data synchronisation", check_text_sync()),
        ("4. Figure outputs and links", check_figure_outputs_and_links()),
        ("5. Guardrails", check_guardrails()),
        ("6. Reference integrity", check_reference_integrity()),
        ("7. SI closure", check_si_closure()),
        ("8. Submission package", check_submission_package()),
    ]
    total_passed = total_failed = total_warned = 0
    for title, checks in sections:
        p, f, w = print_section(title, checks)
        total_passed += p
        total_failed += f
        total_warned += w
    print("\n" + "=" * 72)
    if total_failed:
        print(f"{RED}VALIDATION FAILED{RESET}: {total_failed} hard checks failed, {total_passed} passed")
        if total_warned:
            print(f"{YELLOW}{total_warned} warnings remain{RESET}")
        sys.exit(1)
    print(f"{GREEN}ALL HARD CHECKS PASSED{RESET}: {total_passed} passed")
    if total_warned:
        print(f"{YELLOW}{total_warned} warnings remain{RESET}")
    print("=" * 72)


if __name__ == "__main__":
    main()
