#!/usr/bin/env python3
"""
Generate Word documents for an npj Clean Water Perspective submission.

Converts markdown drafts to formatted .docx files:
  - main_manuscript.docx (Title page, Abstract, Perspective text, back matter,
    References, inline figures and captions)
  - supporting_information.docx (SI title page, Contents, SI body)
  - cover_letter.docx
  - highlights.txt

Style: Times New Roman 12pt, double spacing, A4, 2.54cm margins.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from paths import DRAFTS_DIR, FIGURES_DIR, SUBMISSION_DIR

# =============================================================================
# Paths
# =============================================================================

DRAFTS = DRAFTS_DIR
FIGURES = FIGURES_DIR
OUT_DIR = SUBMISSION_DIR

MAIN_SECTIONS = [
    ("abstract", DRAFTS / "abstract.md"),
    ("introduction", DRAFTS / "introduction.md"),
    ("sec1", DRAFTS / "sec1_landscape.md"),
    ("sec2", DRAFTS / "sec2_data_quality.md"),
    ("sec3", DRAFTS / "sec3_it_ot.md"),
    ("sec4", DRAFTS / "sec4_trust.md"),
    ("sec6", DRAFTS / "sec6_discussion.md"),
    ("data_availability", DRAFTS / "data_availability.md"),
    ("code_availability", DRAFTS / "code_availability.md"),
    ("acknowledgements", DRAFTS / "acknowledgements.md"),
    ("author_contributions", DRAFTS / "author_contributions.md"),
    ("competing_interests", DRAFTS / "competing_interests.md"),
    ("references", DRAFTS / "ref.md"),
]

SI_PATH = DRAFTS / "SI_supporting_information.md"
TITLE_PAGE_PATH = DRAFTS / "Title_page.md"
FIGURE_CAPTIONS_PATH = DRAFTS / "figure_captions.md"
COVER_LETTER_PATH = DRAFTS / "cover_letter.md"
HIGHLIGHTS_PATH = DRAFTS / "highlights.md"

TITLE = "Operational Evidence Standards for Machine Learning in Wastewater Treatment"
TARGET_JOURNAL = "npj Clean Water"
FIGURE_FILES = {
    1: FIGURES / "fig1_landscape.png",
    2: FIGURES / "fig2_operational_illustration.png",
    3: FIGURES / "fig3_evidence_framework.png",
}
FIGURE_HOME_AFTER = {
    "sec2": [1],
    "sec3": [2],
    "sec4": [3],
}

# =============================================================================
# Style constants
# =============================================================================

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
TITLE_SIZE = Pt(12)
HEADING_SIZE = Pt(12)
TABLE_SIZE = Pt(9)
CITATION_BLUE = RGBColor(0x05, 0x63, 0xC1)

# Regex patterns
_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_BOLD_ITALIC_RE = re.compile(r"(\*\*\*(.+?)\*\*\*)|(\*\*(.+?)\*\*)|(\*(.+?)\*)")
_SUBSCRIPT_RE = re.compile(r"([A-Za-z])₂|([A-Za-z])₅₄")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|$")
_BLOCKQUOTE_RE = re.compile(r"^>\s*(.*)$")
_SEE_FIGURE_RE = re.compile(r"^\[See `figures/([^`]+)`\]$")
_IMAGE_FIGURE_RE = re.compile(r"^!\[[^\]]*\]\((?:\.\./\.\./)?figures/([^)]+)\)$")
_CITATION_MARKER_RE = re.compile(r"\[(?:\d+(?:\s*-\s*\d+)?)(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*\]")


# =============================================================================
# Style helpers
# =============================================================================

def set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False, italic=False):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(para, line_spacing=WD_LINE_SPACING.DOUBLE,
                          space_before=Pt(0), space_after=Pt(0)):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after


def setup_document():
    """Create a new document with A4, 2.54cm margins."""
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return doc


def add_line_numbering(doc):
    """Add continuous line numbering to the document."""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement("w:lnNumType")
        lnNumType.set(qn("w:countBy"), "1")
        lnNumType.set(qn("w:restart"), "continuous")
        sectPr.append(lnNumType)


def add_page_numbers(doc):
    """Add centered PAGE fields to every footer."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line_spacing=WD_LINE_SPACING.SINGLE)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), " PAGE ")
        run_el = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)
        rfonts.set(qn("w:eastAsia"), FONT_NAME)
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rpr.append(sz)
        run_el.append(rpr)
        fld.append(run_el)
        p._p.append(fld)


def add_page_break(doc):
    doc.add_page_break()


# =============================================================================
# Inline formatting
# =============================================================================

def _add_formatted_run(para, content, base_size=BODY_SIZE, bold=False,
                       italic=False, citation_blue=False):
    """Add one run and optionally apply direct citation-blue colour."""
    run = para.add_run(content)
    set_font(run, size=base_size, bold=bold, italic=italic)
    if citation_blue:
        run.font.color.rgb = CITATION_BLUE
    return run


def add_runs(para, text, base_size=BODY_SIZE, color_citations=False):
    """Parse markdown inline formatting and add runs to paragraph."""
    # Process bold+italic, bold, italic
    parts = []
    pos = 0
    for m in _BOLD_ITALIC_RE.finditer(text):
        if m.start() > pos:
            parts.append(("normal", text[pos:m.start()]))
        if m.group(1):  # bold+italic
            parts.append(("bold_italic", m.group(2)))
        elif m.group(3):  # bold
            parts.append(("bold", m.group(4)))
        elif m.group(5):  # italic
            parts.append(("italic", m.group(6)))
        pos = m.end()
    if pos < len(text):
        parts.append(("normal", text[pos:]))

    for style, content in parts:
        # Handle markdown links: [text](url) -> just text
        content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
        bold = style in ("bold", "bold_italic")
        italic = style in ("italic", "bold_italic")
        if not color_citations:
            _add_formatted_run(para, content, base_size=base_size, bold=bold, italic=italic)
            continue

        pos = 0
        for match in _CITATION_MARKER_RE.finditer(content):
            if match.start() > pos:
                _add_formatted_run(
                    para,
                    content[pos:match.start()],
                    base_size=base_size,
                    bold=bold,
                    italic=italic,
                )
            _add_formatted_run(
                para,
                match.group(0),
                base_size=base_size,
                bold=bold,
                italic=italic,
                citation_blue=True,
            )
            pos = match.end()
        if pos < len(content):
            _add_formatted_run(
                para,
                content[pos:],
                base_size=base_size,
                bold=bold,
                italic=italic,
            )


# =============================================================================
# Table rendering (three-line / booktabs style)
# =============================================================================

def _set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, start, end with val/sz/color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for key, value in attrs.items():
            el.set(qn(f"w:{key}"), str(value))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _clear_table_borders(table):
    """Remove all borders from a table (prepare for three-line style)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove existing tblBorders if any
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


# Border specs: thick (1.5pt = 12 half-pts), thin (0.75pt = 6 half-pts)
_THICK = {"val": "single", "sz": "12", "color": "000000", "space": "0"}
_THIN = {"val": "single", "sz": "6", "color": "000000", "space": "0"}


def render_table(doc, rows):
    """Render a markdown table as a three-line (booktabs) Word table."""
    if len(rows) < 2:
        return
    # Parse header
    header_cells = [c.strip() for c in rows[0].strip("|").split("|")]
    ncols = len(header_cells)

    # Parse data rows (skip separator)
    data_rows = []
    for row in rows[1:]:
        if _TABLE_SEP_RE.match(row):
            continue
        cells = [c.strip() for c in row.strip("|").split("|")]
        while len(cells) < ncols:
            cells.append("")
        data_rows.append(cells[:ncols])

    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Clear all borders, then apply three-line style
    _clear_table_borders(table)

    # Header row
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, cell_text, base_size=TABLE_SIZE)
        for run in p.runs:
            run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Top thick line + bottom thin line
        _set_cell_border(cell, top=_THICK, bottom=_THIN)

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, cell_text, base_size=TABLE_SIZE)

    # Bottom thick line on last data row
    if data_rows:
        last_row = table.rows[-1]
        for cell in last_row.cells:
            _set_cell_border(cell, bottom=_THICK)

    # Set font for all cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = TABLE_SIZE

    doc.add_paragraph()  # spacing after table


def load_title_page_lines():
    """Load author, affiliation and submission metadata from Title_page.md."""
    if not TITLE_PAGE_PATH.exists():
        return ["[Author names]", "[Affiliations]", f"Submitted to {TARGET_JOURNAL}, [Month Year]"]

    lines = []
    for raw in TITLE_PAGE_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("\\†") and "contributed equally" in line:
            continue
        lines.append(line.replace("\\*", "*").replace("\\†", "†"))
    return lines


def load_title_page_metadata():
    """Parse author, affiliation, contact and submission lines."""
    lines = load_title_page_lines()
    metadata = {
        "authors": "",
        "affiliations": [],
        "contact": [],
        "submitted": f"Submitted to {TARGET_JOURNAL}, June 2026",
    }
    if not lines:
        return metadata

    metadata["authors"] = lines[0]
    for line in lines[1:]:
        if line.startswith("Submitted to"):
            metadata["submitted"] = line
        elif "Corresponding author" in line or line.startswith("E-mail:"):
            metadata["contact"].append(line)
        else:
            metadata["affiliations"].append(line)
    return metadata


def add_metadata_paragraph(doc, text, alignment, bold=False, italic=False,
                           space_before=Pt(0), space_after=Pt(0)):
    """Add a title-page metadata paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    set_paragraph_spacing(p, line_spacing=WD_LINE_SPACING.DOUBLE,
                          space_before=space_before, space_after=space_after)
    run = p.add_run(text)
    set_font(run, size=BODY_SIZE, bold=bold, italic=italic)
    return p


def add_title_page(doc, supporting_information=False):
    """Insert the main or SI title page and end with a page break."""
    metadata = load_title_page_metadata()

    if supporting_information:
        add_metadata_paragraph(
            doc,
            "Supporting Information",
            WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            italic=True,
            space_after=Pt(12),
        )

    add_metadata_paragraph(
        doc,
        TITLE,
        WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        space_after=Pt(18),
    )

    if metadata["authors"]:
        add_metadata_paragraph(doc, metadata["authors"], WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    for line in metadata["affiliations"]:
        add_metadata_paragraph(doc, line, WD_ALIGN_PARAGRAPH.LEFT)

    if metadata["contact"]:
        doc.add_paragraph()
    for line in metadata["contact"]:
        add_metadata_paragraph(doc, line, WD_ALIGN_PARAGRAPH.LEFT)

    add_metadata_paragraph(
        doc,
        metadata["submitted"],
        WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(72),
        space_after=Pt(12),
    )

    add_page_break(doc)


# =============================================================================
# Section processing
# =============================================================================

def _should_color_citations(section_key):
    """Colour main-manuscript citation markers before References only."""
    return section_key not in ("references", "si", "cover_letter")


def process_md_lines(doc, lines, section_key=""):
    """Process markdown lines and add to document."""
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip markdown comments / metadata lines
        if stripped.startswith("> npj Clean Water") or stripped.startswith("> **项目"):
            i += 1
            continue

        # Heading
        m_heading = _HEADING_RE.match(stripped)
        if m_heading:
            level = len(m_heading.group(1))
            heading_text = m_heading.group(2).strip()

            # Skip "# Abstract", "# References", "# Methods" etc as we add our own
            if level == 1 and heading_text.lower() in ("abstract", "references",
                                                        "methods", "figure and table captions"):
                i += 1
                continue

            # For ref.md, skip section headers like "## Introduction", "## sec1 Landscape"
            if section_key == "references" and level >= 2:
                i += 1
                continue

            p = doc.add_paragraph()
            set_paragraph_spacing(p, space_before=Pt(12), space_after=Pt(6))
            if section_key == "cover_letter" and level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(heading_text)
            size = Pt(14) if section_key == "cover_letter" and level == 1 else HEADING_SIZE
            set_font(run, size=size, bold=True)
            i += 1
            continue

        # Table detection
        if _TABLE_ROW_RE.match(stripped):
            table_lines = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            render_table(doc, table_lines)
            continue

        # Blockquote (used in ref.md for notes, SI for prompts)
        if _BLOCKQUOTE_RE.match(stripped):
            quote_text = _BLOCKQUOTE_RE.match(stripped).group(1)
            if quote_text:
                p = doc.add_paragraph()
                set_paragraph_spacing(p)
                p.paragraph_format.left_indent = Cm(1.0)
                add_runs(p, quote_text, color_citations=_should_color_citations(section_key))
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph()
            set_paragraph_spacing(p)
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            add_runs(p, stripped, color_citations=_should_color_citations(section_key))
            i += 1
            continue

        # Horizontal rule
        if stripped.startswith("---"):
            i += 1
            continue

        # SI figure directives: legacy [See `figures/file.png`] or markdown image syntax
        m_see = _SEE_FIGURE_RE.match(stripped)
        m_image = _IMAGE_FIGURE_RE.match(stripped)
        if m_see or m_image:
            filename = m_see.group(1) if m_see else m_image.group(1)
            image_path = FIGURES / filename
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, line_spacing=WD_LINE_SPACING.SINGLE, space_after=Pt(12))
            if image_path.exists():
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(6.0))
            else:
                run = p.add_run(f"[Missing supplementary figure file: {image_path.name}]")
                set_font(run, bold=True)
            i += 1
            continue

        # Reference entries: [1] Author...
        if section_key == "references" and re.match(r"^\[\d+\]", stripped):
            # Only include the reference line, skip annotation lines
            p = doc.add_paragraph()
            set_paragraph_spacing(p)
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            add_runs(p, stripped, color_citations=False)
            i += 1
            # Skip annotation lines (starting with "- ")
            while i < len(lines) and lines[i].strip().startswith("- "):
                i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        set_paragraph_spacing(p)
        p.paragraph_format.first_line_indent = Pt(24)
        add_runs(p, stripped, color_citations=_should_color_citations(section_key))
        i += 1


def parse_figure_captions(text):
    """Return {figure_number: caption_text} from figure_captions.md."""
    captions = {}
    parts = re.split(r"^## Figure\s+(\d+)\s*$", text, flags=re.M)
    for i in range(1, len(parts), 2):
        num = int(parts[i])
        body = parts[i + 1].strip()
        body = re.sub(r"^# Figure Captions\s*", "", body).strip()
        captions[num] = body
    return captions


def add_figure_with_caption(doc, fig_num, image_path, caption):
    """Insert a figure image followed by its caption."""
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=WD_LINE_SPACING.SINGLE, space_after=Pt(6))
    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.0))
    else:
        run = p.add_run(f"[Missing figure file: {image_path.name}]")
        set_font(run, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=WD_LINE_SPACING.DOUBLE, space_after=Pt(12))
    add_runs(p, caption, color_citations=True)


def strip_si_front_matter(lines):
    """Drop SI source cover/title lines; the DOCX builder renders them separately."""
    for idx, line in enumerate(lines):
        if re.match(r"^##\s+S\d+", line.strip()):
            return lines[idx:]
    return lines


def extract_si_contents(lines):
    """Return top-level SI headings for the Contents page."""
    entries = []
    for line in strip_si_front_matter(lines):
        m = re.match(r"^##\s+(S\d+\..+)$", line.strip())
        if m:
            entries.append(m.group(1).strip())
    return entries


def add_si_contents(doc, entries):
    """Insert a one-level SI contents page."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, space_before=Pt(12), space_after=Pt(12))
    run = p.add_run("Contents")
    set_font(run, size=HEADING_SIZE, bold=True)

    for entry in entries:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=WD_LINE_SPACING.DOUBLE)
        add_runs(p, entry, color_citations=False)

    add_page_break(doc)


# =============================================================================
# Main manuscript generation
# =============================================================================

def generate_main_manuscript():
    print("Generating main_manuscript.docx...")
    doc = setup_document()
    add_line_numbering(doc)
    add_page_numbers(doc)
    captions = parse_figure_captions(FIGURE_CAPTIONS_PATH.read_text(encoding="utf-8"))

    add_title_page(doc)

    # Process each section
    for section_key, filepath in MAIN_SECTIONS:
        if not filepath.exists():
            print(f"  WARNING: {filepath} not found, skipping")
            continue

        text = filepath.read_text(encoding="utf-8")
        lines = text.split("\n")

        # Add section headers where needed
        if section_key == "abstract":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, space_before=Pt(12))
            run = p.add_run("Abstract")
            set_font(run, size=HEADING_SIZE, bold=True)

        elif section_key == "introduction":
            add_page_break(doc)

        elif section_key in (
            "data_availability",
            "code_availability",
            "acknowledgements",
            "author_contributions",
            "competing_interests",
        ):
            add_page_break(doc)

        elif section_key == "references":
            add_page_break(doc)
            p = doc.add_paragraph()
            set_paragraph_spacing(p, space_before=Pt(12))
            run = p.add_run("References")
            set_font(run, size=HEADING_SIZE, bold=True)

        process_md_lines(doc, lines, section_key=section_key)

        # Add keywords after abstract
        if section_key == "abstract":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, space_before=Pt(12))
            run = p.add_run("Keywords: ")
            set_font(run, bold=True)
            run = p.add_run("machine learning, water treatment, operational evidence, "
                            "uncertainty quantification, process control")
            set_font(run)

        for fig_num in FIGURE_HOME_AFTER.get(section_key, []):
            add_figure_with_caption(doc, fig_num, FIGURE_FILES.get(fig_num, Path()), captions[fig_num])

    out = OUT_DIR / "main_manuscript.docx"
    doc.save(str(out))
    print(f"  Saved: {out} ({out.stat().st_size / 1024:.1f} KB)")


# =============================================================================
# SI generation
# =============================================================================

def generate_si():
    print("Generating supporting_information.docx...")
    doc = setup_document()
    add_line_numbering(doc)
    add_page_numbers(doc)

    text = SI_PATH.read_text(encoding="utf-8")
    lines = text.split("\n")
    body_lines = strip_si_front_matter(lines)

    add_title_page(doc, supporting_information=True)
    add_si_contents(doc, extract_si_contents(lines))
    process_md_lines(doc, body_lines, section_key="si")

    out = OUT_DIR / "supporting_information.docx"
    doc.save(str(out))
    print(f"  Saved: {out} ({out.stat().st_size / 1024:.1f} KB)")


# =============================================================================
# Cover letter and highlights
# =============================================================================

def generate_cover_letter():
    print("Generating cover_letter.docx...")
    doc = setup_document()
    if COVER_LETTER_PATH.exists():
        lines = COVER_LETTER_PATH.read_text(encoding="utf-8").split("\n")
    else:
        lines = default_cover_letter_lines()
    process_md_lines(doc, lines, section_key="cover_letter")
    out = OUT_DIR / "cover_letter.docx"
    doc.save(str(out))
    print(f"  Saved: {out} ({out.stat().st_size / 1024:.1f} KB)")


def default_cover_letter_lines():
    """Fallback cover letter if the markdown source is missing."""
    metadata = load_title_page_metadata()
    contact = " ".join(metadata["contact"])
    affiliation = metadata["affiliations"][0] if metadata["affiliations"] else ""
    return [
        "# Cover Letter",
        "",
        "June 2026",
        "",
        "Editor",
        "*npj Clean Water*",
        "",
        "Dear Editor,",
        "",
        f"We submit \"{TITLE}\" for consideration as a Perspective in *{TARGET_JOURNAL}*.",
        "",
        "The manuscript addresses how machine-learning studies in water and wastewater treatment should report evidence when claims move from retrospective prediction to operational use. It combines a 423-study full-text evidence map, a bounded six-plant operational illustration and an operation-aware evidence standard covering sensing context, validation timing, uncertainty, control risk and human-ML operating roles.",
        "",
        "We believe the manuscript fits the journal because it connects water and wastewater treatment, process control, data-driven modelling and responsible implementation. The accompanying public package includes the literature corpus, manual-review records, extraction prompts, benchmark summaries, figure source data and anonymised plant-derived outputs. Raw plant records remain confidential under utility agreements, but derived outputs are provided to verify the reported plant-level summaries.",
        "",
        "The manuscript is original, is not under consideration elsewhere and all authors have approved its submission. We do not provide suggested or excluded reviewers at this stage.",
        "",
        "Sincerely,",
        "",
        contact or "Prof. & Dr. Xiuwen Cheng",
        affiliation,
    ]


def generate_highlights():
    print("Generating highlights.txt...")
    if HIGHLIGHTS_PATH.exists():
        text = HIGHLIGHTS_PATH.read_text(encoding="utf-8")
        lines = []
        for raw in text.splitlines():
            stripped = raw.strip()
            if not stripped or stripped.startswith("#"):
                continue
            lines.append(re.sub(r"^-+\s*", "", stripped))
    else:
        lines = [
            "423 full-text studies show sparse operational ML evidence",
            "Six WWTP datasets illustrate routine-monitoring limits",
            "Sensing context and time validation should be reported explicitly",
            "Deployment claims need uncertainty, controls and human-role evidence",
        ]

    out = OUT_DIR / "highlights.txt"
    out.write_text("\n".join(f"- {line}" for line in lines) + "\n", encoding="utf-8")
    print(f"  Saved: {out} ({out.stat().st_size / 1024:.1f} KB)")


# =============================================================================
# Entry point
# =============================================================================

def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_main_manuscript()
    generate_si()
    generate_cover_letter()
    generate_highlights()
    print(f"\nDone! Files saved to: {OUT_DIR}")


if __name__ == "__main__":
    main()
